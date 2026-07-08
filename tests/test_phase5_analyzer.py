import os
import sys
import time
from pathlib import Path

import pytest
from PIL import Image, ImageDraw
from docx import Document
from docx.shared import Inches

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from hermes.agents.analyzer import build_source_analysis, source_analysis_to_json
from hermes.core.docx_ingest import ingest_docx
from hermes.core.llm_config import get_llm
from hermes.core.risk import RISK_MATRIX
from hermes.core.verifier import CHECKER_REGISTRY, verify_artifact
from hermes.rubrics import load_rubric


@pytest.fixture
def sample_docx(tmp_path):
    img1 = tmp_path / 'chart1.png'
    canvas1 = Image.new('RGB', (420, 260), 'white')
    draw1 = ImageDraw.Draw(canvas1)
    draw1.rectangle((60, 100, 150, 220), fill='red')
    draw1.rectangle((220, 60, 310, 220), fill='blue')
    draw1.text((40, 20), 'Chart 1: A=12, B=18', fill='black')
    canvas1.save(img1)

    img2 = tmp_path / 'chart2.png'
    canvas2 = Image.new('RGB', (420, 260), 'white')
    draw2 = ImageDraw.Draw(canvas2)
    draw2.ellipse((60, 60, 190, 190), fill='green')
    draw2.text((220, 120), 'Growth 25%', fill='black')
    canvas2.save(img2)

    docx_path = tmp_path / 'sample.docx'
    doc = Document()
    doc.add_paragraph('Overview paragraph explains the study setup.')
    doc.add_paragraph('First image is referenced immediately before insertion.')
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = 'Metric'
    table.cell(0, 1).text = 'Value'
    table.cell(1, 0).text = 'Accuracy'
    table.cell(1, 1).text = '92%'
    table.cell(2, 0).text = 'Samples'
    table.cell(2, 1).text = '128'
    doc.add_picture(str(img1), width=Inches(4.0))
    doc.add_paragraph('Second image appears after the growth note.')
    doc.add_picture(str(img2), width=Inches(4.0))
    doc.save(docx_path)
    return docx_path


def test_ingest_docx_extracts_paragraphs_tables_images(sample_docx, tmp_path):
    result = ingest_docx(str(sample_docx), str(tmp_path / 'images'))
    assert len(result['paragraphs']) == 3
    assert len(result['tables']) == 1
    assert len(result['images']) == 2
    assert result['tables'][0][1][1] == '92%'


def test_ingest_docx_preserves_image_order(sample_docx, tmp_path):
    result = ingest_docx(str(sample_docx), str(tmp_path / 'images'))
    assert [img['order'] for img in result['images']] == [1, 2]
    assert result['images'][0]['position_context'] == 'First image is referenced immediately before insertion.'
    assert result['images'][1]['position_context'] == 'Second image appears after the growth note.'


def test_local_cx_connectivity():
    required = [
        'HERMES_LOCAL_CX_BASE_URL',
        'HERMES_LOCAL_CX_API_KEY',
        'HERMES_LOCAL_CX_MODEL',
    ]
    if not all(os.environ.get(name) for name in required):
        pytest.skip('local_cx server credentials not configured in environment for live test')

    os.environ['HERMES_LLM_PROVIDER'] = 'local_cx'
    llm = get_llm()
    start = time.time()
    resp = llm.call('Tra loi dung 1 tu: OK')
    elapsed = time.time() - start
    assert 'OK' in resp.upper()
    assert elapsed < 120


def test_analyzer_continues_on_single_image_failure(sample_docx, tmp_path, monkeypatch):
    ingested = ingest_docx(str(sample_docx), str(tmp_path / 'images'))

    class FakeLLM:
        def __init__(self):
            self.calls = 0

        def describe_image(self, image_url, prompt):
            self.calls += 1
            if self.calls == 2:
                raise RuntimeError('vision timeout')
            return f'description-{self.calls}'

    monkeypatch.setattr('hermes.agents.analyzer.get_llm', lambda provider=None: FakeLLM())
    analysis = build_source_analysis(ingested)

    assert len(analysis['images']) == 2
    assert analysis['images'][0]['description'] == 'description-1'
    assert analysis['images'][1]['description'].startswith('[LOI DOC ANH: vision timeout]')
    assert analysis['paragraphs_summary']


def test_source_analysis_registered_in_risk_matrix():
    assert RISK_MATRIX['source_analysis'] == 'low'
    assert 'source_analysis' in CHECKER_REGISTRY


def test_source_analysis_rubric_passes(sample_docx, tmp_path, monkeypatch):
    ingested = ingest_docx(str(sample_docx), str(tmp_path / 'images'))

    class FakeLLM:
        def describe_image(self, image_url, prompt):
            return 'Bar chart with values 12 and 18.'

    monkeypatch.setattr('hermes.agents.analyzer.get_llm', lambda provider=None: FakeLLM())
    analysis = build_source_analysis(ingested)
    rubric = load_rubric('source_analysis')
    result = verify_artifact('source_analysis', source_analysis_to_json(analysis), rubric)
    assert result['passed'] is True
    assert result['detail']['has_content'] == 1.0
    assert result['detail']['images_processed'] == 1.0
