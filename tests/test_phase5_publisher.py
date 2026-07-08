"""Unit tests for Phase 5.5 — Publisher (paper_draft → .docx)."""

import json
import tempfile
from pathlib import Path

import pytest
from docx import Document

from hermes.agents.publisher import publish_paper_draft
from hermes.core.verifier import verify_artifact, CHECKER_REGISTRY
from hermes.rubrics import load_rubric
from hermes.core.risk import should_trigger_debate, SKIP_DEBATE_TYPES


# ── Sample paper (valid IMRaD) ────────────────────────────────────────
SAMPLE_PAPER = """# Evaluation of Model Performance

## Abstract
This study achieved 92% accuracy with 128 samples.

## Introduction
Accuracy is a fundamental metric for classification models.

## Methods
Data were obtained from the source analysis document.

## Results
The model achieved 92% accuracy on 128 samples.

| Metric   | Value |
|----------|-------|
| Accuracy | 92%   |
| Samples  | 128   |

## Discussion
The accuracy indicates strong performance, though sample size is limited.

## References
Data provided by the attached source analysis document.
"""


# ── Tests: Publisher ──────────────────────────────────────────────────

def test_publish_creates_docx_file():
    """publish_paper_draft must create a real .docx file."""
    output = Path(tempfile.mkdtemp()) / "test-output.docx"
    result = publish_paper_draft(SAMPLE_PAPER, output)
    assert result == output
    assert output.is_file()
    assert output.stat().st_size > 0


def test_publish_docx_has_imrad_headings():
    """Generated .docx must contain IMRaD section headings."""
    output = Path(tempfile.mkdtemp()) / "test-imrad.docx"
    publish_paper_draft(SAMPLE_PAPER, output)

    doc = Document(str(output))
    headings = [
        p.text.strip().lower()
        for p in doc.paragraphs
        if p.style.name.startswith("Heading")
    ]
    heading_text = " ".join(headings)

    # Must have at minimum abstract, methods, results
    assert "abstract" in heading_text, f"Missing Abstract. Headings: {headings}"
    assert "method" in heading_text, f"Missing Methods. Headings: {headings}"
    assert "result" in heading_text, f"Missing Results. Headings: {headings}"


def test_publish_docx_contains_table():
    """Markdown tables must be converted to Word tables."""
    output = Path(tempfile.mkdtemp()) / "test-table.docx"
    publish_paper_draft(SAMPLE_PAPER, output)

    doc = Document(str(output))
    tables = doc.tables
    assert len(tables) >= 1, "Expected at least one table in .docx"

    # Verify table content
    table = tables[0]
    assert "Accuracy" in table.cell(1, 0).text
    assert "92%" in table.cell(1, 1).text


def test_publish_handles_minimal_paper():
    """Minimal paper (just a title) should still produce valid .docx."""
    output = Path(tempfile.mkdtemp()) / "test-minimal.docx"
    publish_paper_draft("# Just a Title", output)
    assert output.is_file()
    doc = Document(str(output))
    assert len(doc.paragraphs) > 0


# ── Tests: final_paper checker ────────────────────────────────────────

def test_final_paper_registered_in_checker():
    """final_paper must be registered in CHECKER_REGISTRY."""
    assert "final_paper" in CHECKER_REGISTRY, (
        f"final_paper missing from CHECKER_REGISTRY. "
        f"Keys: {sorted(CHECKER_REGISTRY)}"
    )


def test_final_paper_rubric_pass_with_valid_docx():
    """Valid .docx file must pass the final_paper rubric."""
    output = Path(tempfile.mkdtemp()) / "test-rubric-pass.docx"
    publish_paper_draft(SAMPLE_PAPER, output)

    rubric = load_rubric("final_paper")
    artifact_content = json.dumps({"docx_path": str(output)})

    result = verify_artifact("final_paper", artifact_content, rubric)
    assert result["passed"] is True, (
        f"Expected pass, got score={result['score']}, detail={result['detail']}"
    )
    assert result["score"] >= rubric["pass_threshold"]


def test_final_paper_rubric_fail_missing_file():
    """Non-existent .docx file must fail the rubric."""
    rubric = load_rubric("final_paper")
    artifact_content = json.dumps({"docx_path": "/nonexistent/file.docx"})

    result = verify_artifact("final_paper", artifact_content, rubric)
    assert result["passed"] is False
    assert result["detail"].get("file_exists", 1.0) == 0.0


def test_final_paper_rubric_fail_empty_json():
    """Empty JSON payload must fail."""
    rubric = load_rubric("final_paper")
    result = verify_artifact("final_paper", "{}", rubric)
    assert result["passed"] is False
    assert result["detail"].get("file_exists", 1.0) == 0.0


# ── Tests: skip debate for final_paper ────────────────────────────────

def test_final_paper_in_skip_debate_types():
    """final_paper must NOT trigger a second debate review."""
    assert "final_paper" in SKIP_DEBATE_TYPES, (
        f"final_paper must be in SKIP_DEBATE_TYPES to avoid redundant debate. "
        f"Current: {SKIP_DEBATE_TYPES}"
    )
    assert should_trigger_debate("final_paper") is False, (
        "final_paper is a format conversion of already-reviewed paper_draft. "
        "It must NOT trigger a second debate."
    )
