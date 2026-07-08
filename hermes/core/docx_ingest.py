"""Extract paragraphs, tables, and inline images from a .docx file."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def ingest_docx(file_path: str, image_output_dir: str) -> dict[str, Any]:
    """Return extracted paragraphs, tables, and images from a .docx file."""
    doc = Document(file_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = [_extract_table(table) for table in doc.tables]
    images = _extract_images(doc, image_output_dir)
    return {
        "paragraphs": paragraphs,
        "tables": tables,
        "images": images,
    }


def _extract_table(table: Table) -> list[list[str]]:
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def _extract_images(doc: DocumentType, output_dir: str) -> list[dict[str, Any]]:
    """Save each inline image and preserve its order in the document flow."""
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    images: list[dict[str, Any]] = []
    last_non_empty_paragraph = ""
    sequence = 0

    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                last_non_empty_paragraph = text

            for run in block.runs:
                blips = run._element.findall(f".//{{{A_NS}}}blip")
                for blip in blips:
                    rel_id = blip.get(qn("r:embed"))
                    if not rel_id:
                        continue
                    image_part = doc.part.related_parts[rel_id]
                    extension = _extension_for_content_type(image_part.content_type)
                    sequence += 1
                    filename = f"image_{sequence:03d}{extension}"
                    image_path = output_path / filename
                    image_path.write_bytes(image_part.blob)
                    images.append(
                        {
                            "path": str(image_path),
                            "filename": filename,
                            "content_type": image_part.content_type,
                            "order": sequence,
                            "position_context": text or last_non_empty_paragraph,
                        }
                    )
    return images


def _iter_block_items(doc: DocumentType):
    """Yield Paragraph and Table objects in document order."""
    parent = doc.element.body
    for child in parent.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def _extension_for_content_type(content_type: str) -> str:
    mapping = {
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/jpg": ".jpg",
        "image/gif": ".gif",
        "image/bmp": ".bmp",
        "image/tiff": ".tiff",
    }
    return mapping.get(content_type, ".bin")
