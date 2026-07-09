"""
Rule-based verifier for Hermes artifacts.

Phase 1+2 scope: rubric-based evaluation for 4 artifact types:
  - lit_review_md   (Phase 1)
  - course_outline  (Phase 2)
  - lecture_draft   (Phase 2)
  - quiz_bank       (Phase 2)

All checkers are deterministic, reproducible, and cheap — no LLM judge.
"""

import json
import re
from typing import Dict, Any

from .storage import update_verification
from .risk import get_risk_level, get_effective_threshold
from .events import log_event


# -------------------------------------------------------------------
# Human gate: types that must be manually approved after passing rubric
# -------------------------------------------------------------------
HUMAN_GATE_TYPES = {"lecture_draft", "quiz_bank"}


# -------------------------------------------------------------------
# Checker registry — mirrors the PROVIDER_REGISTRY pattern from llm_config
# -------------------------------------------------------------------
CHECKER_REGISTRY: Dict[str, Any] = {}


def checker_for(artifact_type: str):
    """Decorator to register a checker function for an artifact type."""
    def decorator(fn):
        CHECKER_REGISTRY[artifact_type] = fn
        return fn
    return decorator


def verify_artifact(artifact_type: str, content: str, rubric: dict) -> dict:
    """Dispatch to the correct checker based on artifact_type.

    Raises:
        ValueError: if artifact_type is not registered.
    """
    checker = CHECKER_REGISTRY.get(artifact_type)
    if checker is None:
        registered = ", ".join(sorted(CHECKER_REGISTRY.keys()))
        raise ValueError(
            f"No checker registered for artifact_type '{artifact_type}'. "
            f"Registered types: {registered}"
        )
    return checker(content, rubric)


# -------------------------------------------------------------------
# Centralized verification finalization — single source of truth for
# human-gate policy.  Use this instead of calling update_verification()
# directly so every code path (pipeline, test, CLI) applies the same rules.
# -------------------------------------------------------------------
def finalize_verification(
    workspace,
    artifact_id: str,
    version: int,
    artifact_type: str,
    rubric_result: dict,
    notes: str = "",
    rubric_pass_threshold: float | None = None,
    debate_verdict: dict | None = None,
) -> str:
    """
    Determine and persist the final verification status for an artifact.

    Policy (Phase 3 — risk-adjusted):
      1. Compute risk level and effective threshold from risk matrix.
      2. If rubric_pass_threshold is provided, re-evaluate pass/fail
         using effective_threshold (risk-adjusted floor).
      3. rubric fails  → status = "fail"
      4. rubric passes, type in HUMAN_GATE_TYPES → status = "escalated"
      5. rubric passes, debate_verdict provided:
           - "no_consensus" → "escalated"
           - "consensus_fail" → "fail"
           - "consensus_pass" → "pass" (unless human gate)
      6. rubric passes, otherwise → status = "pass"

    Args:
        workspace:  Workspace instance
        artifact_id: artifact identifier
        version:   artifact version number
        artifact_type: one of the registered artifact types
        rubric_result: dict with "passed" (bool), "score" (float), "detail" (dict)
        notes:     optional human-readable note (e.g. retry count)
        rubric_pass_threshold: base threshold from rubric (for risk adjustment)
        debate_verdict: optional debate review result dict

    Returns:
        The status string that was written to the artifact index.
    """
    risk_level = get_risk_level(artifact_type)

    # ── Risk-adjusted pass/fail ──────────────────────────────────────
    if rubric_pass_threshold is not None:
        effective_threshold = get_effective_threshold(artifact_type, rubric_pass_threshold)
        rubric_pass = rubric_result["score"] >= effective_threshold
    else:
        effective_threshold = rubric_pass_threshold  # None
        rubric_pass = rubric_result["passed"]

    # ── Determine status ─────────────────────────────────────────────
    # Special handling: debate_verdict artifact verified on its own
    # (not passed as debate_verdict parameter to finalize_verification).
    # The final_decision field from the checker overrides rubric score.
    verdict_decision = rubric_result.get("final_decision", "")
    if artifact_type == "debate_verdict" and debate_verdict is None and verdict_decision:
        if verdict_decision == "no_consensus":
            status = "escalated"
        elif verdict_decision == "consensus_fail":
            status = "fail"  # academic rejection — no retry
        else:
            # consensus_pass — fall through to normal rubric scoring
            if not rubric_pass:
                status = "fail"
            else:
                status = "pass"
    elif not rubric_pass:
        status = "fail"
    elif debate_verdict is not None:
        decision = debate_verdict.get("final_decision", "no_consensus")
        if decision == "no_consensus":
            status = "escalated"
        elif decision == "consensus_fail":
            status = "fail"
        elif decision == "consensus_pass":
            if artifact_type in HUMAN_GATE_TYPES:
                status = "escalated"
            else:
                status = "pass"
        else:
            status = "escalated"  # safety: unknown decision → escalate
    elif artifact_type in HUMAN_GATE_TYPES:
        status = "escalated"
    else:
        status = "pass"

    # ── Build verification notes with risk info ──────────────────────
    detail = rubric_result.get("detail", {})
    risk_info = {
        "risk_level": risk_level,
        "effective_threshold": effective_threshold,
        "rubric_score": rubric_result.get("score"),
    }
    if debate_verdict is not None:
        risk_info["debate_final_decision"] = debate_verdict.get("final_decision")
        risk_info["debate_rounds"] = len(debate_verdict.get("rounds", []))

    detail_json = json.dumps(detail, ensure_ascii=False)
    risk_json = json.dumps(risk_info, ensure_ascii=False)
    update_verification(
        workspace,
        artifact_id,
        version,
        status,
        notes=f"{notes}  {detail_json}  risk={risk_json}".strip(),
    )

    # ── Log verification event ───────────────────────────────────────
    log_event(workspace, "verification_result", {
        "artifact_id": artifact_id,
        "artifact_version": version,
        "artifact_type": artifact_type,
        "status": status,
        "risk_level": risk_level,
        "rubric_score": rubric_result.get("score"),
        "effective_threshold": effective_threshold,
    })

    return status


# -------------------------------------------------------------------
# Shared scoring helper
# -------------------------------------------------------------------
def _score(artifact_type: str, content: str, rubric: dict) -> dict:
    """Common return shape for all checkers."""
    checker = CHECKER_REGISTRY.get(artifact_type)
    if checker is None:
        raise ValueError(f"No checker for '{artifact_type}'")
    result = checker(content, rubric)
    return result


# -------------------------------------------------------------------
# Checker: lit_review_md (Phase 1 — kept as-is, now decorated)
# -------------------------------------------------------------------
@checker_for("lit_review_md")
def check_lit_review(content: str, rubric: dict) -> dict:
    scores: Dict[str, float] = {}
    text_lower = content.lower()
    criterion_names = {c["name"] for c in rubric.get("criteria", [])}

    if "citation_completeness" in criterion_names:
        citation_count = len(re.findall(r"\(\d{4}\)|\([A-Za-z\-]+[,\s]+\d{4}\)", content))
        scores["citation_completeness"] = min(citation_count / 3.0, 1.0)

    if "relevance" in criterion_names or "relevance_summary" in criterion_names:
        score = 1.0 if re.search(r"(^|#+\s*)summary", text_lower) else 0.0
        if "relevance" in criterion_names:
            scores["relevance"] = score
        if "relevance_summary" in criterion_names:
            scores["relevance_summary"] = score

    if "gap_identification" in criterion_names or "gaps_section" in criterion_names:
        score = 1.0 if "gaps identified" in text_lower else 0.0
        if "gap_identification" in criterion_names:
            scores["gap_identification"] = score
        if "gaps_section" in criterion_names:
            scores["gaps_section"] = score

    if "clarity" in criterion_names or "formatting_clarity" in criterion_names:
        word_count = len(content.split())
        if word_count > 100:
            score = 1.0
        elif word_count > 50:
            score = 0.5
        else:
            score = 0.0
        if "clarity" in criterion_names:
            scores["clarity"] = score
        if "formatting_clarity" in criterion_names:
            scores["formatting_clarity"] = score

    return _build_result(scores, rubric)


# -------------------------------------------------------------------
# Checker: course_outline (Phase 2)
# -------------------------------------------------------------------
@checker_for("course_outline")
def check_course_outline(content: str, rubric: dict) -> dict:
    scores: Dict[str, float] = {}
    text_lower = content.lower()
    criterion_names = {c["name"] for c in rubric.get("criteria", [])}

    if "learning_objectives_present" in criterion_names:
        has_lo = bool(re.search(r"learning\s+objective", text_lower))
        scores["learning_objectives_present"] = 1.0 if has_lo else 0.0

    if "session_breakdown" in criterion_names:
        has_breakdown = bool(re.search(r"(week|session|tuần|buổi)\s+\d+", text_lower))
        scores["session_breakdown"] = 1.0 if has_breakdown else 0.0

    if "aligned_with_lit_review" in criterion_names:
        has_alignment = bool(re.search(r"\(\d{4}\)", content))
        scores["aligned_with_lit_review"] = 1.0 if has_alignment else 0.0

    if "assessment_hooks" in criterion_names:
        has_assessment = bool(re.search(r"assessment|quiz|exam|kiểm\s*tra", text_lower))
        scores["assessment_hooks"] = 1.0 if has_assessment else 0.0

    return _build_result(scores, rubric)


# -------------------------------------------------------------------
# Checker: lecture_draft (Phase 2)
# -------------------------------------------------------------------
@checker_for("lecture_draft")
def check_lecture_draft(content: str, rubric: dict) -> dict:
    scores: Dict[str, float] = {}
    text_lower = content.lower()
    criterion_names = {c["name"] for c in rubric.get("criteria", [])}

    if "covers_all_outline_sections" in criterion_names:
        word_count = len(content.split())
        scores["covers_all_outline_sections"] = 1.0 if word_count > 300 else 0.0

    if "examples_included" in criterion_names:
        has_example = bool(re.search(r"for example|ví dụ|example|instance", text_lower))
        scores["examples_included"] = 1.0 if has_example else 0.0

    if "length_adequate" in criterion_names:
        word_count = len(content.split())
        scores["length_adequate"] = 1.0 if word_count > 500 else 0.0

    if "no_unsupported_claims" in criterion_names:
        bare_numbers = len(re.findall(r"(?<!\()\b\d{4,}(?!\))", content))
        citation_markers = len(re.findall(r"\(\d{4}\)", content))
        scores["no_unsupported_claims"] = 0.0 if bare_numbers > 0 and citation_markers == 0 else 1.0

    return _build_result(scores, rubric)


# -------------------------------------------------------------------
# Checker: quiz_bank (Phase 2)
# -------------------------------------------------------------------
@checker_for("quiz_bank")
def check_quiz_bank(content: str, rubric: dict) -> dict:
    scores: Dict[str, float] = {}
    text_lower = content.lower()
    criterion_names = {c["name"] for c in rubric.get("criteria", [])}

    if "question_count" in criterion_names:
        # Match numbered questions at start of a line.
        # Patterns: "1. ", "2) ", "Q1. ", "A. ", "**Q1.** ", "** 1.** ", "- Q1) "
        question_lines = re.findall(
            r"^\s*(?:\d+[.)]\s|(?:Q?\d+|\*\*)[.)]\s|(?:^|[^-])[A-Z]\.\s)",
            content,
            re.MULTILINE,
        )
        q_count = len(question_lines)
        scores["question_count"] = 1.0 if q_count >= 5 else q_count / 5.0

    if "covers_lecture_topics" in criterion_names:
        has_topic_markers = bool(re.search(r"(lecture|chủ đề|topic|bài học)", text_lower))
        scores["covers_lecture_topics"] = 1.0 if has_topic_markers else 0.0

    if "has_answer_key" in criterion_names:
        has_answer = bool(re.search(r"(?:đáp án|answer|correct answer|=>\s*[A-Z])", text_lower))
        scores["has_answer_key"] = 1.0 if has_answer else 0.0

    if "difficulty_variety" in criterion_names:
        easy_markers = re.findall(r"\b(easy|dễ|beginner|basic)\b", text_lower)
        hard_markers = re.findall(r"\b(hard|khó|advanced|expert)\b", text_lower)
        has_variety = bool(easy_markers) and bool(hard_markers)
        mc_markers = len(re.findall(r"\n\s*[A-Z][.)]\s", content))
        tf_markers = len(re.findall(r"\b(true|false|đúng|sai)\b", text_lower))
        has_variety = has_variety or (mc_markers > 0 and tf_markers > 0)
        scores["difficulty_variety"] = 1.0 if has_variety else 0.0

    return _build_result(scores, rubric)


# -------------------------------------------------------------------
# Checker: paper_draft (Phase 5.2) — rule-based only; LLM-judge in P5.3
# -------------------------------------------------------------------
@checker_for("paper_draft")
def check_paper_draft(content: str, rubric: dict) -> dict:
    """Check paper_draft structure against IMRaD rubric.

    Rule-based only.  ``data_fidelity`` and ``reviewer_verdict`` scores
    are defaulted — the LLM-judge in P5.3 will replace them with real
    assessments.
    """
    import re
    scores: Dict[str, float] = {}
    criterion_names = {c["name"] for c in rubric.get("criteria", [])}
    text_lower = content.lower()

    if "structure_completeness" in criterion_names:
        required_sections = [
            ("title", [r"title", r"ti[uù] đề"]),
            ("abstract", [r"abstract", r"t[oó]m t[ắa]t"]),
            ("introduction", [r"introduction", r"gi[ớơ]i thi[ệe]u", r"m[ởơ] đ[ầa]u"]),
            ("methods", [r"methods?", r"ph[uư][ơo]ng ph[aá]p"]),
            ("results", [r"results?", r"k[ếe]t qu[ảa]"]),
            ("discussion", [r"discussion", r"th[ảa]o lu[ậa]n", r"b[aà]n lu[ậa]n"]),
            ("references", [r"references?", r"t[aà]i li[ệe]u tham kh[ảa]o"]),
        ]
        found = 0
        for _name, patterns in required_sections:
            for pat in patterns:
                if re.search(rf"(?:^|\n)#+\s*{pat}", text_lower):
                    found += 1
                    break
        scores["structure_completeness"] = found / len(required_sections)

    if "data_fidelity" in criterion_names:
        # Cannot verify without comparing against source_analysis input.
        # P5.3 LLM-judge will handle this.  Default: neutral-passing.
        scores["data_fidelity"] = 0.7

    if "citation_format" in criterion_names:
        has_refs = bool(re.search(r"(?:references|tài liệu tham khảo)", text_lower))
        entries = len(re.findall(r"\(\d{4}\)", content))
        if has_refs and entries >= 3:
            scores["citation_format"] = 1.0
        elif has_refs:
            # Check if the References section has the anti-fabrication fallback
            # instead of real citations — that deserves full credit too.
            if re.search(
                r"data provided by|nguồn dữ liệu được cung cấp|không có tài liệu tham khảo",
                text_lower,
            ):
                scores["citation_format"] = 1.0
            else:
                scores["citation_format"] = 0.5
        else:
            scores["citation_format"] = 0.0

    if "clarity" in criterion_names:
        word_count = len(content.split())
        scores["clarity"] = 1.0 if word_count > 300 else (
            0.5 if word_count > 150 else 0.0
        )

    if "reviewer_verdict" in criterion_names:
        # LLM-judge will provide real verdict in P5.3.
        scores["reviewer_verdict"] = 1.0

    return _build_result(scores, rubric)


# -------------------------------------------------------------------
# Checker: debate_verdict (Phase 3)
# -------------------------------------------------------------------
@checker_for("source_analysis")
def check_source_analysis(content: str, rubric: dict) -> dict:
    """Check source_analysis artifact coverage without any LLM judge."""
    scores: Dict[str, float] = {}
    criterion_names = {c["name"] for c in rubric.get("criteria", [])}

    try:
        payload = json.loads(content) if isinstance(content, str) else content
    except json.JSONDecodeError:
        payload = {}

    if "has_content" in criterion_names:
        summary = str(payload.get("paragraphs_summary", "")).strip()
        scores["has_content"] = 1.0 if summary else 0.0

    if "images_processed" in criterion_names:
        images = payload.get("images", []) if isinstance(payload, dict) else []
        all_described = bool(images) and all(str(img.get("description", "")).strip() for img in images)
        scores["images_processed"] = 1.0 if all_described else 0.0

    return _build_result(scores, rubric)


# -------------------------------------------------------------------
# Checker: literature_support (Phase 5.2.5)
# -------------------------------------------------------------------
@checker_for("literature_support")
def check_literature_support(content: str, rubric: dict) -> dict:
    """Check literature_support entries have real url + excerpt from API."""
    scores: Dict[str, float] = {}
    criterion_names = {c["name"] for c in rubric.get("criteria", [])}

    try:
        payload = json.loads(content) if isinstance(content, str) else content
    except json.JSONDecodeError:
        payload = {}

    entries = payload.get("entries", []) if isinstance(payload, dict) else []
    search_attempted = payload.get("search_attempted", False) if isinstance(payload, dict) else False
    queries = payload.get("queries_used", []) if isinstance(payload, dict) else []
    search_error = payload.get("search_error") if isinstance(payload, dict) else None

    if "has_url_for_all" in criterion_names:
        if not entries:
            scores["has_url_for_all"] = 1.0  # empty = trivially satisfied
        else:
            all_have_url = all(
                isinstance(e, dict) and str(e.get("url", "")).strip()
                for e in entries
            )
            scores["has_url_for_all"] = 1.0 if all_have_url else 0.0

    if "has_excerpt_for_all" in criterion_names:
        if not entries:
            scores["has_excerpt_for_all"] = 1.0
        else:
            all_have_excerpt = all(
                isinstance(e, dict) and str(e.get("excerpt", "")).strip()
                for e in entries
            )
            scores["has_excerpt_for_all"] = 1.0 if all_have_excerpt else 0.0

    if "search_attempted" in criterion_names:
        # search_error only matters when we got zero results from all queries.
        # Partial rate-limiting that still returned entries is acceptable.
        if search_error and not entries:
            scores["search_attempted"] = 0.0
        else:
            scores["search_attempted"] = 1.0 if (
                search_attempted and isinstance(queries, list) and len(queries) > 0
            ) else 0.0

    return _build_result(scores, rubric)


# -------------------------------------------------------------------
# Checker: debate_verdict (Phase 3)
# -------------------------------------------------------------------
@checker_for("debate_verdict")
def check_debate_verdict(content: str, rubric: dict) -> dict:
    """Check debate_verdict artifact against its rubric.

    The content should be a JSON string or dict representation of the verdict.
    """
    scores: Dict[str, float] = {}
    criterion_names = {c["name"] for c in rubric.get("criteria", [])}

    # Try to parse as JSON
    verdict = None
    if isinstance(content, str):
        try:
            verdict = json.loads(content)
        except json.JSONDecodeError:
            verdict = {}
    elif isinstance(content, dict):
        verdict = content
    else:
        verdict = {}

    if "consensus_reached" in criterion_names:
        decision = verdict.get("final_decision", "")
        has_consensus = decision in ("consensus_pass", "consensus_fail")
        scores["consensus_reached"] = 1.0 if has_consensus else 0.0

    if "rounds_completed" in criterion_names:
        rounds = verdict.get("rounds", [])
        scores["rounds_completed"] = 1.0 if len(rounds) >= 1 else 0.0

    if "arguments_present" in criterion_names:
        rounds = verdict.get("rounds", [])
        all_have_args = all(
            r.get("proponent_argument", "").strip() and r.get("opponent_argument", "").strip()
            for r in rounds
        )
        scores["arguments_present"] = 1.0 if all_have_args and len(rounds) > 0 else 0.0

    return _build_result(scores, rubric, extra={"final_decision": verdict.get("final_decision", "")})


# -------------------------------------------------------------------
# Rule-based diff guard: Editor must not change numbers or add fake citations
# (Phase 5.6 — called BEFORE Reviewer LLM-judge on retry attempts)
# -------------------------------------------------------------------
def check_editor_diff(
    original: str,
    edited: str,
    source_analysis: dict[str, Any],
    literature_support: dict[str, Any] | None = None,
) -> tuple[bool, list[str]]:
    """Verify Editor only changed what it was allowed to change.

    Returns (all_valid, list_of_violations).

    Rules:
      1. Every NUMBER in the edited version must exist in the original
         OR in source_analysis.key_statistics.
      2. Every NEW citation (not in original) must match an entry in
         literature_support.
    """
    import re
    violations: list[str] = []

    # ── Rule 1: number guard ────────────────────────────────────────
    allowed_numbers = set()
    for ks in source_analysis.get("key_statistics", []) or []:
        val = str(ks).strip().rstrip("%")
        allowed_numbers.add(val)
        allowed_numbers.add(val + "%")  # also allow with percent sign
    for m in re.finditer(r"\b\d+(?:\.\d+)?%?\b", original):
        allowed_numbers.add(m.group(0))

    edited_numbers = set(re.findall(r"\b\d+(?:\.\d+)?%?\b", edited))
    # Exclude years (4-digit 19xx/20xx) — they belong to citations, not data
    new_numbers = {
        n for n in (edited_numbers - allowed_numbers)
        if not (len(n) == 4 and n.startswith(("19", "20")))
    }
    if new_numbers:
        violations.append(
            f"Editor introduced {len(new_numbers)} new numbers not in original "
            f"or source: {sorted(new_numbers)[:5]}"
        )

    # ── Rule 2: citation guard ──────────────────────────────────────
    if literature_support:
        entries = literature_support.get("entries", []) or []
        if entries:
            lit_authors = set()
            for e in entries:
                authors = (e.get("authors", "") or "").lower()
                for author in authors.split(","):
                    lit_authors.add(author.strip())

            orig_citations = set(
                m.group(0).lower()
                for m in re.finditer(
                    r"[A-Z][a-z]+(?:\s*,\s*[A-Z]\.)?\s*\(\s*\d{4}\s*\)",
                    original,
                )
            )
            edited_citations = set(
                m.group(0).lower()
                for m in re.finditer(
                    r"[A-Z][a-z]+(?:\s*,\s*[A-Z]\.)?\s*\(\s*\d{4}\s*\)",
                    edited,
                )
            )
            new_citations = edited_citations - orig_citations

            for cite in new_citations:
                author_part = cite.split("(")[0].strip().lower()
                found = any(author_part in la for la in lit_authors)
                if not found:
                    found = any(
                        word in cite
                        for la in lit_authors
                        for word in la.split()
                        if len(word) > 2
                    )
                if not found:
                    violations.append(
                        f"Editor added citation '{cite}' not in literature_support"
                    )

    # ── Rule 3: citation year change guard ──────────────────────────
    # For citations that exist in BOTH versions with the same author
    # but a different year, verify the new year against literature_support.
    if literature_support:
        entries = literature_support.get("entries", []) or []
        if entries:
            # Build index: (author_lower, year) → True
            lit_index = {}
            for e in entries:
                authors = (e.get("authors", "") or "").lower()
                year = str(e.get("year", "") or "")
                for author in authors.split(","):
                    author_key = author.strip()
                    if author_key and year:
                        lit_index[(author_key, year)] = True

            # Extract (author, year) pairs from both versions
            cite_pattern = re.compile(
                r"([A-Z][a-z]+(?:\s*,\s*[A-Z]\.)?)\s*\(\s*(\d{4})\s*\)",
                re.IGNORECASE,
            )
            orig_pairs = {(m.group(1).strip().lower(), m.group(2)) for m in cite_pattern.finditer(original)}
            edited_pairs = {(m.group(1).strip().lower(), m.group(2)) for m in cite_pattern.finditer(edited)}

            for author, year in edited_pairs:
                # Find same author in original with different year
                orig_years = {y for a, y in orig_pairs if a == author}
                if orig_years and year not in orig_years:
                    # Year changed — verify against literature
                    if (author, year) not in lit_index:
                        # Also check without year specificity (author exists in any year)
                        author_has_any_year = any(
                            a == author for a, y in lit_index
                        )
                        if not author_has_any_year:
                            violations.append(
                                f"Editor changed citation year: '{author}' "
                                f"from {orig_years} to {year}, "
                                f"not found in literature_support"
                            )

    return len(violations) == 0, violations


# -------------------------------------------------------------------
# Checker: final_paper (Phase 5.5) — rule-based: file exists + not empty + IMRaD
# -------------------------------------------------------------------
@checker_for("final_paper")
def check_final_paper(content: str, rubric: dict) -> dict:
    """Check final_paper artifact: .docx file exists, not empty, has IMRaD."""
    import json
    from pathlib import Path
    from docx import Document

    scores: Dict[str, float] = {}
    criterion_names = {c["name"] for c in rubric.get("criteria", [])}

    # Parse content (JSON with file path)
    try:
        payload = json.loads(content) if isinstance(content, str) else content
    except json.JSONDecodeError:
        payload = {}
    docx_path = payload.get("docx_path", "")

    if "file_exists" in criterion_names:
        scores["file_exists"] = 1.0 if docx_path and Path(docx_path).is_file() else 0.0

    if "file_not_empty" in criterion_names:
        if docx_path and Path(docx_path).is_file():
            try:
                doc = Document(docx_path)
                has_content = len(doc.paragraphs) > 0
                scores["file_not_empty"] = 1.0 if has_content else 0.0
            except Exception:
                scores["file_not_empty"] = 0.0
        else:
            scores["file_not_empty"] = 0.0

    if "has_imrad_headings" in criterion_names:
        if docx_path and Path(docx_path).is_file():
            try:
                doc = Document(docx_path)
                heading_texts = [
                    p.text.strip().lower()
                    for p in doc.paragraphs
                    if p.style.name.startswith("Heading")
                ]
                imrad_keywords = [
                    "abstract", "introduction", "method",
                    "result", "discussion", "reference",
                ]
                found = sum(
                    1 for kw in imrad_keywords
                    if any(kw in h for h in heading_texts)
                )
                scores["has_imrad_headings"] = found / len(imrad_keywords)
            except Exception:
                scores["has_imrad_headings"] = 0.0
        else:
            scores["has_imrad_headings"] = 0.0

    return _build_result(scores, rubric)


# -------------------------------------------------------------------
# Shared result builder
# -------------------------------------------------------------------
def _build_result(scores: Dict[str, float], rubric: dict, extra: dict | None = None) -> dict:
    """Compute weighted total and build the standard result dict."""
    weighted_total = sum(
        scores.get(c["name"], 0.0) * c["weight"]
        for c in rubric.get("criteria", [])
    )
    threshold = rubric.get("pass_threshold", 0.7)
    passed = weighted_total >= threshold

    result = {
        "passed": passed,
        "score": round(weighted_total, 3),
        "detail": scores,
    }
    if extra:
        result.update(extra)
    return result
