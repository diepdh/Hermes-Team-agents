"""Publisher subagent for Phase 5.5.

Converts a verified paper_draft (markdown) into a .docx file with
proper IMRaD formatting.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _set_heading(doc: Document, text: str, level: int = 1) -> None:
    """Add a heading paragraph with consistent styling."""
    p = doc.add_heading(text, level=level)
    return p


def _add_body(doc: Document, text: str) -> None:
    """Add a body paragraph."""
    if text.strip():
        doc.add_paragraph(text.strip())


def publish_paper_draft(
    paper_draft_content: str,
    output_path: str | Path,
    title: str | None = None,
) -> Path:
    """Convert a paper_draft (IMRaD markdown) into a .docx file.

    Args:
        paper_draft_content: The markdown content from paper_draft artifact.
        output_path: Where to write the .docx file (absolute or relative).
        title: Optional override for the document title.

    Returns:
        Path to the generated .docx file.

    The function parses markdown headings (#, ##, ###) to identify IMRaD
    sections and converts them to Word headings.  Body paragraphs between
    headings become Word paragraphs.  Tables (markdown pipe syntax) are
    converted to Word tables.
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # ── Parse markdown into sections ──────────────────────────────────
    lines = paper_draft_content.split("\n")
    sections: list[dict[str, Any]] = []  # {level, heading, body_lines}
    current_section: dict[str, Any] | None = None

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("#"):
            # Count heading level
            level = 0
            for ch in stripped:
                if ch == "#":
                    level += 1
                else:
                    break
            heading_text = stripped[level:].strip()
            if current_section is not None:
                sections.append(current_section)
            current_section = {
                "level": min(level, 3),  # max level 3
                "heading": heading_text,
                "body_lines": [],
            }
        elif current_section is not None:
            current_section["body_lines"].append(line)

    if current_section is not None:
        sections.append(current_section)

    # ── Build .docx ───────────────────────────────────────────────────
    for sec in sections:
        heading = sec["heading"]
        body_lines = sec["body_lines"]

        # Title → centered heading 1
        if sec["level"] == 1:
            _set_heading(doc, heading, level=1)
        else:
            _set_heading(doc, heading, level=sec["level"])

        # Parse body: split into paragraphs, detect tables
        body_text = "\n".join(body_lines)
        paragraphs = body_text.split("\n\n")

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            # Detect markdown tables (lines starting/ending with |)
            table_lines = [l.strip() for l in para.split("\n") if l.strip().startswith("|")]
            if len(table_lines) >= 2:
                # Try to parse as table
                rows = []
                for tl in table_lines:
                    cells = [c.strip() for c in tl.strip("|").split("|")]
                    # Skip separator rows (like |---|---|)
                    if all(c.replace("-", "").replace(":", "").strip() == "" for c in cells):
                        continue
                    if cells:
                        rows.append(cells)

                if len(rows) >= 1:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = "Light Grid Accent 1"
                    for i, row_cells in enumerate(rows):
                        for j, cell_text in enumerate(row_cells):
                            if j < len(rows[0]):
                                table.cell(i, j).text = cell_text
                    doc.add_paragraph()  # spacing after table
                    continue

            # Regular paragraph
            _add_body(doc, para)

    # ── Set document properties ───────────────────────────────────────
    if title:
        doc.core_properties.title = title

    # ── Save ──────────────────────────────────────────────────────────
    doc.save(str(output_path))
    return output_path
